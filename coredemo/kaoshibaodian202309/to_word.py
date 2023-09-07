import json

import pymysql
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

db = pymysql.connect(host="localhost",
                     port=3306,
                     user="root",
                     password="root",
                     db="bigdata",
                     charset="utf8")
cursor = db.cursor()
table_name = 'bigdata.`test_record`'

query_sql = "select * from %s order by id asc " % (table_name)



per_record = {
    "StyleID": 16,
    "Type": "ATEST",
    "Style": "A1型题",
    "SubType": "",
    "Score": 1,
    "AllTestID": 8433,
    "ChildTableID": -1,
    "SrcID": 1,
    "SbjID": 5,
    "CptID": 244,
    "Title": "d38d3961d5185987779d62fefebdafa691d05639bdf77ea9cf2f93dd4922714e80c7a94edc1d98e705a89c35022e5cbf067acde086b2d5bb1be691be06a11bed7add3f4dc3237c62",

    "Answer": "D",
    "Explain": "百白破混合剂应在第3个月注射，故不选A;麻疹疫苗第一次是在儿童8个月的时候，不选B;乙肝第二次在出生后第1个月，不选C;卡介苗出生后即注射，E也不选。正确答案为D。解题关键：按照我国卫生部的规定，婴儿须在出生后的第2、3、4个月分别注射脊髓灰质炎三价混合疫苗。",
    "TestPoint": "儿童计划免疫",
    "DealInfo": "",
    "OperateTime": "2023年5月05日",
    "pointData": [],
    "knowPoint": "",
    "material_content": "",
    "chapter_content": "",
    "isLNZT": 0,
    "IsFav": 0,
    "UserNoteContent": "",
    "SelectedItems": [
        {
            "HasImg": 1,
            "ItemName": "A",
            "Content": "百白破第一次"
        },
        {
            "HasImg": 0,
            "ItemName": "B",
            "Content": "麻疹第一次"
        },
        {
            "HasImg": 0,
            "ItemName": "C",
            "Content": "乙肝第二次"
        },
        {
            "HasImg": 0,
            "ItemName": "D",
            "Content": "脊髓灰质炎第一次"
        },
        {
            "HasImg": 0,
            "ItemName": "E",
            "Content": "卡介苗"
        }
    ],
    "ReplyRate": 0.5584,
    "LevelRate": 0.4416
}

def sql_to_word():
    # 创建一个新的 Word 文档
    doc = Document()

    # 设置标题
    title = doc.add_heading('考试宝典', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 创建字体样式
    font = doc.styles['Normal'].font
    font.name = '宋体'  # 设置字体为宋体
    font.size = Pt(14)  # 设置字体大小为14磅

    # JSON 数据示例
    json_data1 = {'name': 'John', 'age': 30, 'city': 'New York'}
    json_data2 = {'name': 'Alice', 'age': 25, 'city': 'Los Angeles'}

    per_record = {
        "StyleID": 16,
        "Type": "ATEST",
        "Style": "A1型题",
        "SubType": "",
        "Score": 1,
        "AllTestID": 8433,
        "ChildTableID": -1,
        "SrcID": 1,
        "SbjID": 5,
        "CptID": 244,
        "Title": "d38d3961d5185987779d62fefebdafa691d05639bdf77ea9cf2f93dd4922714e80c7a94edc1d98e705a89c35022e5cbf067acde086b2d5bb1be691be06a11bed7add3f4dc3237c62",
        "Answer": "D",
        "Explain": "百白破混合剂应在第3个月注射，故不选A;麻疹疫苗第一次是在儿童8个月的时候，不选B;乙肝第二次在出生后第1个月，不选C;卡介苗出生后即注射，E也不选。正确答案为D。解题关键：按照我国卫生部的规定，婴儿须在出生后的第2、3、4个月分别注射脊髓灰质炎三价混合疫苗。",
        "TestPoint": "儿童计划免疫",
        "DealInfo": "",
        "OperateTime": "2023年5月05日",
        "pointData": [],
        "knowPoint": "",
        "material_content": "",
        "chapter_content": "",
        "isLNZT": 0,
        "IsFav": 0,
        "UserNoteContent": "",
        "SelectedItems": [
            {
                "HasImg": 1,
                "ItemName": "A",
                "Content": "百白破第一次"
            },
            {
                "HasImg": 0,
                "ItemName": "B",
                "Content": "麻疹第一次"
            },
            {
                "HasImg": 0,
                "ItemName": "C",
                "Content": "乙肝第二次"
            },
            {
                "HasImg": 0,
                "ItemName": "D",
                "Content": "脊髓灰质炎第一次"
            },
            {
                "HasImg": 0,
                "ItemName": "E",
                "Content": "卡介苗"
            }
        ],
        "ReplyRate": 0.5584,
        "LevelRate": 0.4416
    }

    result_length = cursor.execute(query_sql)

    print("result_length: ", result_length)
    db.commit()

    # 获取查询结果
    result = cursor.fetchall()
    print("result: ", result)

    # 将 JSON 数据写入文档
    count = 1
    for data in result:
        print("data: ", data)

        # 添加一个段落
        para = doc.add_paragraph()
        all_test_id = data[10]
        print("all_test_id: ", all_test_id)
        style = data[7]
        cpt_id = data[14]
        title_de = data[16]

        # 将 JSON 数据转换为字符串，并设置段落内容
        #json_str = json.dumps(data, ensure_ascii=False)  # 确保汉字以原始形式写入

        json_str = f"第{cpt_id}章-题号{all_test_id}-{style}：\n {count}. {title_de}"
        run = para.add_run(json_str)
        # 设置段落对齐方式为左对齐
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 添加第二段 针对有选项的处理方法
        para2 = doc.add_paragraph()
        selected_items_str = data[29]
        if selected_items_str != "":
            # 字符串转py对象
            selected_items = json.loads(selected_items_str)
            print("selected_items: ", selected_items)
            selected_items_len = len(selected_items)
            print("selected_items_len: ", selected_items_len)
            for i in range(len(selected_items)):
                print("i : ", i)
                ItemName = selected_items[i]['ItemName']
                Content = selected_items[i]['Content']
                json_str2 = f"{ItemName} {Content} \n"
                run = para2.add_run(json_str2)
                # 设置段落对齐方式为左对齐
                para2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 添加第三段  答案 解析
        para3  = doc.add_paragraph()
        answer = data[17]
        json_str3 = f"答案： {answer} \n"
        run = para3.add_run(json_str3)

        test_explain = data[18]
        json_str4 = f"解析： {test_explain} "
        run = para3.add_run(json_str4)

        run = para3.add_run('\n')
        count +=1

    # 保存文档
    doc.save('ksbd_y_20230906.docx')

if __name__ == "__main__":
    print("start")

    sql_to_word()